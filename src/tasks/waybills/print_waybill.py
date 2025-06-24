from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from schemas.waybill_offer_schema import WaybillOfferSchema
from schemas.waybill_schema import WaybillSchema

templates_path = Path("templates")


def set_cell_margins(cell, left, right, top, bottom):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = tcPr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcPr.append(margins)
    for margin, value in [
        ("left", left),
        ("right", right),
        ("top", top),
        ("bottom", bottom),
    ]:
        element = OxmlElement(f"w:{margin}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        margins.append(element)


def set_cell_borders(cell, top=0, left=0, bottom=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border, val in [
        ("top", top),
        ("left", left),
        ("bottom", bottom),
        ("right", right),
    ]:
        border_elem = OxmlElement(f"w:{border}")
        if val == 1:
            border_elem.set(qn("w:val"), "single")
            border_elem.set(qn("w:sz"), "4")
            border_elem.set(qn("w:space"), "0")
            border_elem.set(qn("w:color"), "auto")
        else:
            border_elem.set(qn("w:val"), "nil")
        tcBorders.append(border_elem)
    tcPr.append(tcBorders)


def create_document(
    waybill: WaybillSchema, offers: list[WaybillOfferSchema], total_sum: float
) -> Path:
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"

    # Заголовок
    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_par.add_run(f"Накладная № {waybill.id}")
    run.bold = True
    run.font.size = Pt(12)

    # Таблица №1 — данные клиента
    table1 = doc.add_table(rows=4, cols=2)
    table1.autofit = False
    table1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    widths1 = [Cm(2.0), Cm(14.0)]
    table1_data = [
        ("ФИО:", waybill.author or "Не указано"),
        ("Телефон:", "Не указан"),
        ("Город:", "Не указан"),
        ("Доставка:", "Самовывоз"),
    ]

    for i, (label, value) in enumerate(table1_data):
        for j, text in enumerate([label, value]):
            cell = table1.rows[i].cells[j]
            cell.width = widths1[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 50 if j == 0 else 0, 50, 30, 30)
            set_cell_borders(cell, 0, 0, 0, 0)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.clear()
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.bold = j == 0

    # Таблица №2 — товары
    table2 = doc.add_table(rows=len(offers) + 2, cols=8)
    table2.autofit = True
    table2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    widths2 = [Cm(1.2), Cm(4.5), Cm(6.0), Cm(3.0), Cm(3.0), Cm(2.0), Cm(2.5), Cm(3.0)]
    for row in table2.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths2[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 50, 50, 50, 50)

    headers = ["#", "Категория", "", "Бренд", "Номер", "Цена", "Кол-во", "Сум."]
    for j, cell in enumerate(table2.rows[0].cells):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(headers[j])
        run.font.size = Pt(9)
        run.bold = True
        set_cell_borders(cell, 1, 1, 1, 1)

    # Данные по каждой позиции
    for i, offer in enumerate(offers):
        row = table2.rows[i + 1].cells
        row[0].text = str(i + 1)
        row[1].paragraphs[0].add_run(offer.category_name).font.size = Pt(7)
        row[2].paragraphs[0].add_run(offer.product_name).font.size = Pt(7.5)
        row[3].paragraphs[0].add_run(offer.brand).font.size = Pt(7.5)
        row[4].paragraphs[0].add_run(offer.manufacturer_number).font.size = Pt(7.5)
        row[5].paragraphs[0].add_run(f"{offer.price_rub:.0f} руб.").font.size = Pt(7.5)
        row[6].paragraphs[0].add_run(str(offer.quantity)).font.size = Pt(7.5)
        sum_price = offer.price_rub * offer.quantity
        row[7].paragraphs[0].add_run(f"{sum_price:.0f} руб.").font.size = Pt(7.5)

        for cell in row:
            set_cell_borders(cell, 0, 0, 1, 0)

    # Последняя строка — итого
    last_row = table2.rows[-1]
    for cell in last_row.cells:
        cell.text = ""
        set_cell_borders(cell, 0, 0, 1, 0)

    last_row.cells[1].paragraphs[0].add_run("Итого").bold = True
    total_run = last_row.cells[7].paragraphs[0].add_run(f"{total_sum:.0f} руб.")
    total_run.bold = True
    total_run.font.size = Pt(7.5)

    # Сохранение
    folder = Path("tmp/")
    folder.mkdir(parents=True, exist_ok=True)
    file_path = folder / f"waybill_{waybill.id}.docx"
    doc.save(str(file_path))

    return file_path
