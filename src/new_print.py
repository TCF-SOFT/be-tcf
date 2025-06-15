from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from io import BytesIO
import datetime

TABLE_STYLE = {
    "border_color": "000000",
    "border_size": 12,
    "header_fill": "DCDCDC",
    "cell_padding": Inches(0.05)
}
FONT_STYLE = {
    "name": "Times New Roman",
    "header_size": Pt(16),
    "body_size": Pt(12),
    "table_header_size": Pt(11)
}
DATE_FORMAT = "%I:%M %p CEST, %d %B %Y"

class InvoiceGenerator:
    def __init__(self):
        self.doc = Document()
        self.doc.styles['Normal'].font.name = FONT_STYLE["name"]

    def _set_cell_border(self, cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for border_type in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_type}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(TABLE_STYLE["border_size"]))
            border.set(qn('w:color'), TABLE_STYLE["border_color"])
            tcBorders.append(border)

    def _set_cell_shading(self, cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

    def _apply_table_style(self, table):
        table.autofit = False
        widths = [Inches(0.5), Inches(2.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(0.7), Inches(1.0)]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                self._set_cell_border(cell)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = FONT_STYLE["name"]
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_STYLE["name"])

        for cell in table.rows[0].cells:
            self._set_cell_shading(cell, TABLE_STYLE["header_fill"])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_STYLE["table_header_size"]
                    run.bold = True

    def _validate_data(self, data):
        required_keys = {"number", "fio", "phone", "city", "delivery", "items", "total"}
        if not all(key in data for key in required_keys):
            raise ValueError("Отсутствуют обязательные поля в данных")
        if not data["items"] or not all(isinstance(item, dict) and all(k in item for k in ["name", "brand", "number", "price", "quantity"]) for item in data["items"]):
            raise ValueError("Некорректные данные в списке товаров")

    def _calculate_total(self, items):
        return sum(item["price"] * item["quantity"] for item in items)

    def create_invoice(self, data, output_mode="file", filename="invoice.docx") -> str | BytesIO:
        try:
            self._validate_data(data)
            calc_total = self._calculate_total(data["items"])
            if data["total"] != calc_total:
                print(f"Предупреждение: Итоговая сумма не совпадает. Использую {calc_total}")
                data["total"] = calc_total

            header = self.doc.add_paragraph()
            run = header.add_run(f"НАКЛАДНАЯ № {data['number']}")
            run.font.size = FONT_STYLE["header_size"]
            run.bold = True
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.doc.add_paragraph(f"ФИО: {data['fio']}")
            self.doc.add_paragraph(f"Телефон: {data['phone']}")
            self.doc.add_paragraph(f"Город: {data['city']}")
            self.doc.add_paragraph(f"Доставка: {data['delivery']}")
            self.doc.add_paragraph(f"Дата: {datetime.datetime.now().strftime(DATE_FORMAT)}")

            table = self.doc.add_table(rows=1, cols=7)
            headers = ["№", "Наименование", "Бренд", "Номер", "Цена", "Кол-во", "Сумма"]
            for i, text in enumerate(headers):
                table.cell(0, i).text = text

            for idx, item in enumerate(data["items"], 1):
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = item["name"]
                row[2].text = item["brand"]
                row[3].text = item["number"]
                row[4].text = f"{item['price']} руб."
                row[5].text = str(item["quantity"])
                row[6].text = f"{item['price'] * item['quantity']} руб."

            self._apply_table_style(table)

            total = self.doc.add_paragraph()
            total_run = total.add_run(f"Итого: {data['total']} руб.")
            total_run.bold = True
            total_run.font.size = FONT_STYLE["body_size"]

            if output_mode == "file":
                self.doc.save(filename)
                return filename
            else:
                stream = BytesIO()
                self.doc.save(stream)
                stream.seek(0)
                return stream

        except Exception as e:
            raise Exception(f"Ошибка при создании накладной: {str(e)}")

invoice_data = {
    "number": "51240 13793277",
    "fio": "2 Менеджер",
    "phone": "+7 (999) 042-46-66",
    "city": "Не заполнен у пользователя",
    "delivery": "Самовывоз",
    "items": [
        {"name": "Прокладки клапанной крышки", "brand": "Transit BSG", "number": "AR893", "price": 500, "quantity": 10},
        {"name": "Прокладки клапанной крышки", "brand": "Transit пробка", "number": "AR806", "price": 300, "quantity": 5}
    ],
    "total": 9999
}

generator = InvoiceGenerator()
try:
    result = generator.create_invoice(invoice_data, output_mode="file", filename="invoice.docx")
    print(result)
except Exception as e:
    print("Ошибка:", e)